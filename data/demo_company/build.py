"""
Helix Outdoor — multi-format demo data builder.

Generates:
  pdf/returns_policy_v2_1.pdf
  pdf/saigon_texco_msa.pdf
  pdf/q2_vendor_risk_report.pdf
  docx/security_access_policy.docx
  docx/sla_matrix.docx
  images/architecture.png
  images/qc_whiteboard.png
  images/org_chart.png

Run from repo root:
  ./src/python_backend/venv/bin/python data/demo_company/build.py
"""

import os
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from PIL import Image, ImageDraw, ImageFont

ROOT = os.path.dirname(os.path.abspath(__file__))
PDF_DIR = os.path.join(ROOT, "pdf")
DOCX_DIR = os.path.join(ROOT, "docx")
IMG_DIR = os.path.join(ROOT, "images")
for d in (PDF_DIR, DOCX_DIR, IMG_DIR):
    os.makedirs(d, exist_ok=True)


# ── Shared styles for PDFs ────────────────────────────────────────────────────
styles = getSampleStyleSheet()
H1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=18, spaceAfter=12, textColor=colors.HexColor("#0F172A"))
H2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, spaceAfter=8, textColor=colors.HexColor("#1E293B"))
H3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11, spaceAfter=4, textColor=colors.HexColor("#334155"))
BODY = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10, leading=14, spaceAfter=6)
SMALL = ParagraphStyle("Small", parent=styles["BodyText"], fontSize=8, leading=11, textColor=colors.HexColor("#64748B"))


def _para(text, style=BODY):
    return Paragraph(text.replace("\n", "<br/>"), style)


# ══════════════════════════════════════════════════════════════════════════════
# PDF 1 — Returns Policy v2.1
# ══════════════════════════════════════════════════════════════════════════════
def build_returns_policy():
    out = os.path.join(PDF_DIR, "returns_policy_v2_1.pdf")
    doc = SimpleDocTemplate(out, pagesize=LETTER,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    flow = []

    flow += [
        _para("Helix Outdoor — Returns &amp; Refunds Policy", H1),
        _para("Version 2.1 · Effective March 15, 2026", SMALL),
        _para("Owner: Camille Rousseau (Head of Customer Care) · Approved by: Bruno Castelli (Finance), Priya Iyer (Legal)", SMALL),
        Spacer(1, 0.2*inch),

        _para("1. Standard Return Window (Physical Products)", H2),
        _para(
            "Helix Outdoor accepts returns of unused physical product within <b>30 days</b> of the delivery date "
            "for a full refund to the original payment method. Lightly-used product within the same window may be "
            "exchanged or refunded as 75% store credit at the discretion of the Customer Care team."
        ),
        _para(
            "Returns must be initiated through our self-service portal at helixoutdoor.com/returns. The system "
            "auto-generates a prepaid USPS or DHL label depending on origin country."
        ),

        _para("2. Subscription Products (digital, including the Trail Box)", H2),
        _para(
            "Subscription cancellations within the first <b>7 days</b> of purchase qualify for a full refund. "
            "After 7 days the subscription is non-refundable but can be cancelled to prevent future renewals."
        ),
        _para(
            "Note (April 2026 update, see legal addendum): EU customers have a 14-day statutory withdrawal "
            "right that overrides the above. See section 4.2 of this document and Priya Iyer's legal note dated "
            "April 30, 2026."
        ),

        _para("3. Refund Authority", H2),
        _para(
            "<b>Important:</b> The authority limits in this section are superseded by Bruno Castelli's "
            "memo of April 28, 2026. The new limits take effect May 1, 2026 and reduce CX-Lead authority "
            "from $2,000 to $1,000. Please refer to that memo for the current limits."
        ),
        _para(
            "(Historical record, kept for audit completeness:) Prior to May 1, 2026: CX Specialist up to "
            "$300 without approval; CX Lead up to $2,000; Camille up to $10,000; Bruno up to $50,000."
        ),

        _para("4. Special Cases", H2),
        _para("4.1 Defective Product", H3),
        _para(
            "Confirmed manufacturing defects (failed seams, broken hardware, defective zippers etc.) qualify "
            "for full refund regardless of return window, plus a 20% goodwill discount code on the next purchase. "
            "Document the defect with a photograph and log a Lot Number in Gorgias."
        ),
        _para("4.2 EU Customers — Digital Products", H3),
        _para(
            "Per EU Directive 2011/83/EU, EU customers (identified by Stripe billing country) have a 14-day "
            "automatic right of withdrawal on digital products. This <b>overrides</b> the 7-day rule in section 2 "
            "for EU customers only."
        ),
        _para("4.3 Custom &amp; Embroidered Orders", H3),
        _para(
            "Custom and embroidered orders are non-returnable except in the case of confirmed manufacturing defect."
        ),
        _para("4.4 Wholesale Accounts (Camp Cosmos, TrekRight, etc.)", H3),
        _para(
            "Wholesale accounts return per the terms in their Master Distribution Agreement. Default: 60 days from "
            "shipment date for unused product, no restocking fee on first return per quarter. Escalations route to "
            "Cara Bennett (CX-B2B Lead)."
        ),

        _para("5. Logging", H2),
        _para(
            "All returns must be logged in Gorgias with the original order ID in the ticket title and the SKU "
            "and lot number in the body. Refunds executed via Stripe Dashboard: Customers → [Customer ID] → "
            "Payments → Refund. Always send the CX-RETURN-CONFIRM template to the customer after processing."
        ),

        _para("6. Crypto and Non-Standard Payment Refunds", H2),
        _para(
            "Helix Outdoor does not currently accept cryptocurrency. If a refund request is received for a payment "
            "made via a non-standard rail (gift cards from a wholesale partner, store credit from a retired SKU, "
            "etc.), escalate to Bruno Castelli and Priya Iyer before processing."
        ),

        _para("Changelog", H2),
        _para(
            "v2.1 — March 15, 2026: Updated wholesale account language; reduced custom-order exemption window. "
            "v2.0 — November 1, 2025: Migration from manual finance-team review to self-service portal. "
            "v1.0 — Pre-2025: All refunds processed manually via finance@helixoutdoor.com."
        ),
    ]

    doc.build(flow)
    print(f"  ✓ {out}")


# ══════════════════════════════════════════════════════════════════════════════
# PDF 2 — Saigon TexCo Master Supply Agreement
# ══════════════════════════════════════════════════════════════════════════════
def build_msa():
    out = os.path.join(PDF_DIR, "saigon_texco_msa.pdf")
    doc = SimpleDocTemplate(out, pagesize=LETTER,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    flow = []

    flow += [
        _para("Master Supply Agreement", H1),
        _para("Helix Outdoor LLC and Saigon TexCo Manufacturing JSC", H3),
        _para("Effective Date: October 1, 2024 · Renewal: November 1, 2026", SMALL),
        _para(
            "Helix Outdoor LLC, a Delaware limited liability company (\"Buyer\"), and Saigon TexCo "
            "Manufacturing Joint Stock Company, a corporation organised under the laws of Vietnam (\"Supplier\"), "
            "agree as follows:",
            BODY,
        ),
        Spacer(1, 0.15*inch),

        _para("1. Scope of Goods", H2),
        _para(
            "Supplier shall manufacture finished outdoor goods to Buyer specification, including the Trailspire "
            "and Crestcamp pack lines, the Ridgeline series, the Foothill apparel line, and any future SKUs the "
            "parties agree to in writing."
        ),

        _para("2. Pricing &amp; Volume Tiers", H2),
    ]

    pricing_data = [
        ["SKU Family", "Tier 1 (1–999u)", "Tier 2 (1k–4,999u)", "Tier 3 (5k+)"],
        ["Trailspire 40L", "$32.40", "$28.10", "$25.60"],
        ["Crestcamp 60L", "$41.20", "$36.00", "$32.80"],
        ["Ridgeline series", "$24.80", "$22.10", "$19.95"],
        ["Foothill hoodie", "$14.60", "$12.40", "$10.85"],
        ["Custom / embroidered uplift", "+12%", "+10%", "+8%"],
    ]
    pricing_table = Table(pricing_data, colWidths=[2.0*inch, 1.4*inch, 1.4*inch, 1.4*inch])
    pricing_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
        ("TOPPADDING", (0, 0), (-1, 0), 6),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
    ]))
    flow.append(pricing_table)
    flow.append(Spacer(1, 0.15*inch))

    flow += [
        _para("3. Lead Times", H2),
        _para(
            "Standard lead time from confirmed Purchase Order to FOB Ho Chi Minh City handover: <b>21 calendar "
            "days</b>. Custom or embroidered orders: add seven (7) calendar days. Supplier shall notify Buyer in "
            "writing within 48 hours of any deviation greater than five (5) days from this commitment."
        ),
        _para(
            "<i>Note (operational): the parties acknowledge that the 21-day standard has not been achievable since "
            "the Tet 2026 holiday period due to capacity sharing arrangements at Supplier's Ho Chi Minh facility. "
            "Buyer-Supplier ops are operating under a verbal interim agreement of 28 days for nylon SKUs and 42 "
            "days for cotton SKUs pending the November 2026 renewal — see Marco Ferraro / Tran Hai operational "
            "memo of May 5, 2026.</i>"
        ),

        _para("4. Payment Terms", H2),
        _para(
            "Net-45 from the date of FOB handover, payable in USD by wire to Supplier's designated account at "
            "Vietcombank. Late payment interest: 1.0% per month."
        ),

        _para("5. Quality Control", H2),
        _para(
            "Supplier agrees to perform pre-ship QC per the Helix QC checklist (Appendix A — see referenced "
            "QC whiteboard photo, May 2026). At minimum: zipper test (10 cycles), seam stitching density "
            "audit, logo placement audit (within 2mm tolerance), lot tag verification."
        ),
        _para(
            "Supplier shall not substitute sub-suppliers (zipper, hardware, fabric) without 30 days written "
            "notice to Buyer's VP of Operations or designated successor."
        ),

        _para("6. Termination &amp; Renewal", H2),
        _para(
            "Initial term: two (2) years from Effective Date, with automatic renewal for successive one (1) "
            "year terms unless either party gives 90 days written notice prior to renewal. Either party may "
            "terminate for material breach with 30 days cure period."
        ),

        _para("7. Confidentiality &amp; IP", H2),
        _para(
            "All Helix Outdoor designs, patterns, tech packs, and marketing materials remain the sole property "
            "of Buyer. Supplier may not disclose, reproduce, or repurpose Buyer IP for any third party."
        ),

        Spacer(1, 0.3*inch),
        _para("Signed:", H3),
        _para("Maya Lin, VP Operations, Helix Outdoor LLC · October 1, 2024", BODY),
        _para("Tran Hai, Account Director, Saigon TexCo Manufacturing JSC · October 1, 2024", BODY),
    ]

    doc.build(flow)
    print(f"  ✓ {out}")


# ══════════════════════════════════════════════════════════════════════════════
# PDF 3 — Q2 2026 Vendor Risk Report
# ══════════════════════════════════════════════════════════════════════════════
def build_vendor_risk():
    out = os.path.join(PDF_DIR, "q2_vendor_risk_report.pdf")
    doc = SimpleDocTemplate(out, pagesize=LETTER,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    flow = []

    flow += [
        _para("Q2 2026 Vendor Risk Report", H1),
        _para("Prepared by: Bruno Castelli, Head of Finance · Reviewed by: Diego Marin, CEO", SMALL),
        Spacer(1, 0.15*inch),

        _para("Executive Summary", H2),
        _para(
            "This report covers the top 8 vendors by spend and flags concentration, contractual, and continuity "
            "risk as of May 1, 2026. <b>The top finding: Saigon TexCo represents 70% of finished-goods spend "
            "with no qualified secondary supplier of comparable capacity. This is up from 55% one year ago.</b>"
        ),
    ]

    table_data = [
        ["Vendor", "Category", "FY26 Spend (est.)", "Risk", "Owner"],
        ["Saigon TexCo", "Manufacturing (primary)", "$3.4M", "HIGH (concentration)", "Marco Ferraro (acting)"],
        ["Bandung Crafted", "Manufacturing (secondary)", "$1.1M", "MEDIUM (no signed MSA)", "Marco Ferraro (acting)"],
        ["ShipBob", "3PL", "$680k", "LOW", "Nia Okafor"],
        ["Stripe", "Payments", "1.9% TPV (~$310k)", "LOW", "Bruno Castelli"],
        ["Klaviyo", "Email / SMS", "$48k", "LOW", "Lena Park"],
        ["Gorgias", "Helpdesk", "$26k", "LOW", "Camille Rousseau"],
        ["Prime Forwarders", "Customs broker", "$92k", "MEDIUM (single-broker)", "Nia Okafor"],
        ["AMD Developer Cloud", "ML / GPU", "$31k (forecast)", "UNKNOWN — no documented owner", "[GAP]"],
    ]
    risk_table = Table(table_data, colWidths=[1.4*inch, 1.5*inch, 1.2*inch, 1.6*inch, 1.5*inch])
    risk_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
        ("TOPPADDING", (0, 0), (-1, 0), 6),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
        ("BACKGROUND", (3, 1), (3, 1), colors.HexColor("#FEE2E2")),
        ("BACKGROUND", (4, 8), (4, 8), colors.HexColor("#FEE2E2")),
    ]))
    flow.append(risk_table)
    flow.append(Spacer(1, 0.2*inch))

    flow += [
        _para("Detailed Findings", H2),

        _para("1. Saigon TexCo concentration risk", H3),
        _para(
            "70% of finished-goods spend with one supplier is unhealthy. The Tet 2026 capacity-share situation "
            "(see operational memo from Marco Ferraro, May 5) demonstrates that we cannot rely on contracted "
            "lead times when their priority customers consolidate. Recommendation: qualify a second nylon-capable "
            "supplier in calendar 2026. NOTE: this work is currently <b>paused</b> due to the VP Ops vacancy."
        ),

        _para("2. Bandung Crafted: no signed MSA", H3),
        _para(
            "Operating under verbal terms set up by Maya Lin (now departed) on a per-PO Net-30 / FOB Jakarta basis. "
            "Priya Iyer is drafting a formal MSA. Until signed, treat as month-to-month relationship with no "
            "indemnification, no IP protection, no QC commitments."
        ),

        _para("3. AMD Developer Cloud: no documented owner", H3),
        _para(
            "We have an active GPU spend on AMD Developer Cloud (BrainOS prototyping, ML pipeline training) but "
            "the vendor account ownership is unclear. Tomas Becker raised this in #general on May 4, 2026. "
            "Action: Sai Krishnan to formally claim ownership and document in the Vendor Registry by May 15."
        ),

        _para("Open questions for the brain to resolve", H2),
        _para(
            "1. Who is the AMD Developer Cloud account owner today? "
            "2. Are we enrolled in the AMD AI Developer Program? Required for hackathon credits. "
            "3. Has the Bandung Crafted MSA been signed yet? "
            "4. What is the formal lead-time agreement with Saigon TexCo as of the November renewal?"
        ),
    ]

    doc.build(flow)
    print(f"  ✓ {out}")


# ══════════════════════════════════════════════════════════════════════════════
# DOCX 1 — Security & Access Policy
# ══════════════════════════════════════════════════════════════════════════════
def build_security_policy():
    out = os.path.join(DOCX_DIR, "security_access_policy.docx")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading("Helix Outdoor — Security & Access Policy", level=0)
    p = doc.add_paragraph()
    p.add_run("Version 1.3 · Effective May 1, 2026 · Owner: Sai Krishnan (CTO) · Approved by: Diego Marin (CEO)").italic = True

    doc.add_heading("1. Credential Management", level=1)
    doc.add_paragraph(
        "All system credentials must be stored in the Helix Outdoor 1Password Business account. Credentials must "
        "never be shared via Slack, email, Notion, code repositories, or any other channel. The 1Password vaults are:"
    )
    doc.add_paragraph("• Production-Infra — Sai Krishnan, Tomas Becker, Jin Ahn", style="List Bullet")
    doc.add_paragraph("• Production-Payments — Bruno Castelli, Liam Walsh, Camille Rousseau", style="List Bullet")
    doc.add_paragraph("• Production-Logistics — Marco Ferraro, Nia Okafor", style="List Bullet")
    doc.add_paragraph("• Marketing-Tools — Lena Park, marketing team", style="List Bullet")

    doc.add_heading("2. Multi-Factor Authentication", level=1)
    doc.add_paragraph(
        "MFA is mandatory on every system that supports it. Hardware keys (YubiKey) are required for: "
        "AWS root, GitHub admin, Stripe, 1Password admin, Shopify admin, Odoo admin. TOTP is acceptable "
        "for everything else."
    )

    doc.add_heading("3. Production Access", level=1)
    doc.add_paragraph(
        "Production access is granted on a need-to-know basis with quarterly review. Engineering staff access "
        "is approved by Sai Krishnan; finance access is approved by Bruno Castelli; logistics access is "
        "approved by the acting Ops Lead (Marco Ferraro until VP Ops is hired)."
    )
    doc.add_paragraph(
        "Production database operations require the standard change-management process: PR review by at least "
        "two engineers, staging dry-run with confirmation, and a written rollback plan attached to the PR. "
        "Production migrations may only run during the Tuesday 02:00–04:00 UTC maintenance window, except for "
        "P0 hotfixes."
    )

    doc.add_heading("4. Webhook Endpoints", level=1)
    doc.add_paragraph(
        "All inbound webhook endpoints (Stripe, Shopify, ShipBob, Klaviyo, Gorgias) must verify a signature "
        "before processing the payload. A missing signature header MUST be treated as a hard failure (400 "
        "response) and a CRITICAL alert to PagerDuty — NOT silently dropped. Signature verification logic "
        "may not be wrapped in a try/except that suppresses failures."
    )
    doc.add_paragraph(
        "Two-engineer review is required for any change to a webhook signature verification path."
    )

    doc.add_heading("5. Vendor Access", level=1)
    doc.add_paragraph(
        "Third-party vendors granted system access are reviewed quarterly. All vendor access is logged in the "
        "Vendor Registry (Notion → Operations → Vendor Access Log). Access is granted for a maximum of 90 "
        "days at a time and renewed on review."
    )

    doc.add_heading("6. Incident Reporting", level=1)
    doc.add_paragraph(
        "Suspected credential compromise: page Sai Krishnan immediately via PagerDuty (sai-security). Do not "
        "attempt remediation alone. Containment (rotate / revoke) precedes investigation. Post a notice in "
        "#security within 30 minutes of detection."
    )

    doc.add_heading("7. Departure Protocol", level=1)
    doc.add_paragraph(
        "On the last day of any departing employee: 1Password access revoked by Sai Krishnan; SSO disabled in "
        "Google Workspace; Slack access revoked; GitHub access revoked; Stripe / Odoo / Shopify access revoked "
        "by the relevant owner. Hardware keys (YubiKey) returned to Tomas Becker. A departure checklist is "
        "tracked per-person in the People Ops Notion page."
    )
    doc.add_paragraph(
        "Note: the departure protocol was followed for Maya Lin on April 12, 2026 — but the supplier-relationship "
        "knowledge transfer was incomplete and is being reconstructed by Marco Ferraro (see ops memo dated "
        "April 15, 2026)."
    )

    doc.save(out)
    print(f"  ✓ {out}")


# ══════════════════════════════════════════════════════════════════════════════
# DOCX 2 — SLA Matrix
# ══════════════════════════════════════════════════════════════════════════════
def build_sla_matrix():
    out = os.path.join(DOCX_DIR, "sla_matrix.docx")
    doc = Document()

    doc.add_heading("Helix Outdoor — Customer Care SLA Matrix", level=0)
    p = doc.add_paragraph()
    p.add_run("Version 2.0 · Effective May 1, 2026 · Owner: Camille Rousseau · Approved by: Diego Marin").italic = True

    doc.add_heading("1. Tier Definitions", level=1)
    doc.add_paragraph(
        "Helix Outdoor customers are classified into four tiers for support purposes:"
    )
    doc.add_paragraph("• DTC Standard — direct-to-consumer customers, all channels, default tier", style="List Bullet")
    doc.add_paragraph("• Helix Trail Club — VIP loyalty (5+ orders rolling 12 months), ~1,400 customers as of May 2026", style="List Bullet")
    doc.add_paragraph("• Wholesale — accounts with a signed Master Distribution Agreement", style="List Bullet")
    doc.add_paragraph("• Strategic Wholesale — top 3 wholesale accounts by revenue (Camp Cosmos, TrekRight, REI test pilot)", style="List Bullet")

    doc.add_heading("2. Response-Time SLAs", level=1)

    table = doc.add_table(rows=5, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Tier", "First Response", "Resolution Target", "Channels"

    rows = [
        ("DTC Standard", "1 business day", "5 business days", "Email + chat"),
        ("Helix Trail Club", "4 business hours", "2 business days", "Email + chat + priority queue"),
        ("Wholesale", "2 business hours", "1 business day", "Email + chat + dedicated AM"),
        ("Strategic Wholesale", "1 hour 24/7", "Same day", "Email + chat + Slack channel + Cara Bennett direct"),
    ]
    for i, r in enumerate(rows, start=1):
        for j, v in enumerate(r):
            table.rows[i].cells[j].text = v

    doc.add_heading("3. Return Windows", level=1)
    doc.add_paragraph(
        "Return windows by customer tier and product type. The longest applicable window always wins for "
        "the customer."
    )
    doc.add_paragraph(
        "• DTC Standard, physical product: 30 days from delivery (per Returns Policy v2.1)\n"
        "• DTC Standard, digital subscription: 7 days non-EU; 14 days EU (per Priya's legal note, April 30 2026)\n"
        "• Helix Trail Club, physical product: 60 days (per Camille's note, May 1 2026)\n"
        "• Wholesale: 60 days from shipment, no restocking fee on first return per quarter\n"
        "• Defective lot (e.g. Lot 24-A-118 zipper): unlimited window, full refund, no return required"
    )

    doc.add_heading("4. Escalation Paths", level=1)
    doc.add_paragraph(
        "DTC tickets: Tier 1 (Petra / Idris) → Tier 2 (Aria Velasquez) → Camille Rousseau."
    )
    doc.add_paragraph(
        "Wholesale tickets: Cara Bennett (CX-B2B Lead) directly. Strategic Wholesale escalations also route "
        "to Diego Marin."
    )
    doc.add_paragraph(
        "Specific escalation rules currently in force:"
    )
    doc.add_paragraph(
        "• Camp Cosmos (Brian Wallace): direct line to Cara Bennett, escalation to Diego Marin if revenue impact "
        "exceeds $10,000 in a single incident.\n"
        "• Any wholesale customer threatening churn: immediate notify Diego, Camille, Bruno.\n"
        "• Any defective-lot complaint: immediate notify Cara Bennett AND Marco Ferraro (factory liaison)."
    )

    doc.add_heading("5. Currently Assigned Customer Success Managers", level=1)
    doc.add_paragraph("• Camp Cosmos — Cara Bennett")
    doc.add_paragraph("• TrekRight Co — Cara Bennett")
    doc.add_paragraph("• REI test pilot — Cara Bennett (in negotiation; lead agreement target Q3)")

    doc.save(out)
    print(f"  ✓ {out}")


# ══════════════════════════════════════════════════════════════════════════════
# Image helpers
# ══════════════════════════════════════════════════════════════════════════════
def _font(size, bold=False):
    candidates = [
        f"/System/Library/Fonts/Supplemental/Arial{' Bold' if bold else ''}.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/System/Library/Fonts/HelveticaNeue.ttc",
        "/Library/Fonts/Arial.ttf",
    ]
    for c in candidates:
        if os.path.exists(c):
            try:
                return ImageFont.truetype(c, size)
            except Exception:
                continue
    return ImageFont.load_default()


def _box(draw, xy, label, sublabel=None, fill=(248, 250, 252), border=(15, 23, 42)):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=border, width=2)
    f = _font(15, bold=True)
    fs = _font(11)
    cx = (x1 + x2) // 2
    cy = (y1 + y2) // 2
    if sublabel:
        draw.text((cx, cy - 12), label, fill=border, font=f, anchor="mm")
        draw.text((cx, cy + 10), sublabel, fill=(71, 85, 105), font=fs, anchor="mm")
    else:
        draw.text((cx, cy), label, fill=border, font=f, anchor="mm")


def _arrow(draw, p1, p2, label=None, color=(15, 23, 42)):
    draw.line([p1, p2], fill=color, width=2)
    import math
    ang = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
    s = 9
    head = [
        (p2[0] - s * math.cos(ang - math.radians(20)), p2[1] - s * math.sin(ang - math.radians(20))),
        p2,
        (p2[0] - s * math.cos(ang + math.radians(20)), p2[1] - s * math.sin(ang + math.radians(20))),
    ]
    draw.line(head, fill=color, width=2)
    if label:
        mx = (p1[0] + p2[0]) // 2
        my = (p1[1] + p2[1]) // 2
        f = _font(11)
        bbox = draw.textbbox((0, 0), label, font=f)
        w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
        pad = 4
        draw.rectangle([mx - w//2 - pad, my - h//2 - pad, mx + w//2 + pad, my + h//2 + pad],
                       fill=(255, 255, 255), outline=(203, 213, 225))
        draw.text((mx, my), label, fill=(30, 41, 59), font=f, anchor="mm")


# ══════════════════════════════════════════════════════════════════════════════
# IMAGE 1 — Architecture diagram
# ══════════════════════════════════════════════════════════════════════════════
def build_architecture_png():
    out = os.path.join(IMG_DIR, "architecture.png")
    W, H = 1400, 800
    img = Image.new("RGB", (W, H), (255, 255, 255))
    d = ImageDraw.Draw(img)

    title = _font(22, bold=True)
    sub = _font(13)
    d.text((40, 28), "Helix Outdoor — Order & Inventory Architecture", fill=(15, 23, 42), font=title)
    d.text((40, 60), "Owner: Sai Krishnan (CTO) · Last updated May 2026 by Jin Ahn", fill=(100, 116, 139), font=sub)

    # nodes
    _box(d, (60, 130, 280, 220), "Shopify", "Storefront / DTC orders",
         fill=(220, 252, 231), border=(22, 101, 52))
    _box(d, (60, 280, 280, 370), "Klaviyo", "Email / SMS",
         fill=(254, 240, 138), border=(133, 77, 14))
    _box(d, (60, 430, 280, 520), "Gorgias", "Helpdesk / tickets",
         fill=(254, 226, 226), border=(127, 29, 29))

    _box(d, (520, 230, 820, 360),
         "Helix Backend (FastAPI)",
         "shopify-svc · erp-bridge · webhooks-svc",
         fill=(224, 231, 255), border=(49, 46, 129))

    _box(d, (1000, 130, 1340, 220), "Odoo (ERP)", "Inventory · POs · Finance",
         fill=(255, 237, 213), border=(154, 52, 18))
    _box(d, (1000, 280, 1340, 370), "ShipBob (3PL)", "Fulfillment + warehouse",
         fill=(207, 250, 254), border=(14, 116, 144))
    _box(d, (1000, 430, 1340, 520), "Saigon TexCo", "Primary factory · Vietnam",
         fill=(252, 231, 243), border=(157, 23, 77))
    _box(d, (1000, 580, 1340, 670), "Bandung Crafted", "Secondary factory · Indonesia",
         fill=(252, 231, 243), border=(157, 23, 77))

    _box(d, (520, 580, 820, 670), "Stripe", "Payments + refunds",
         fill=(220, 252, 231), border=(22, 101, 52))

    # edges
    _arrow(d, (280, 175), (520, 270), "orders webhook")
    _arrow(d, (280, 325), (520, 305), "events stream")
    _arrow(d, (280, 475), (520, 340), "ticket sync")

    _arrow(d, (820, 280), (1000, 175), "PO writes")
    _arrow(d, (820, 305), (1000, 325), "fulfillment events")
    _arrow(d, (820, 340), (1000, 475), "PO submits")
    _arrow(d, (820, 360), (1000, 625), "PO submits")

    _arrow(d, (1340, 325), (1340, 175), "stock updates")  # ShipBob → Odoo (the bug surface)
    _arrow(d, (1000, 475), (820, 360), "ASN")              # factory back to backend
    _arrow(d, (670, 580), (670, 360), "charge / refund")

    # Annotations / gotcha callouts
    note_font = _font(11)
    d.text((830, 310),
           "⚠ ShipBob → Backend webhook MUST verify signature.\n"
           "    Missing X-Shipbob-Signature → 400 + PagerDuty alert.",
           fill=(159, 18, 57), font=note_font)
    d.text((1040, 100), "Owner: Marco Ferraro (acting Ops)", fill=(100, 116, 139), font=note_font)
    d.text((1010, 540), "Owner: Marco Ferraro (was Maya Lin)", fill=(100, 116, 139), font=note_font)
    d.text((40, 760), "Legend: arrows = direction of write/event. Boxes coloured by domain. "
                      "Two-engineer review required on any webhook handler.",
           fill=(71, 85, 105), font=note_font)

    img.save(out, "PNG", optimize=True)
    print(f"  ✓ {out}")


# ══════════════════════════════════════════════════════════════════════════════
# IMAGE 2 — QC Whiteboard photo (synthetic, looks like a phone snap)
# ══════════════════════════════════════════════════════════════════════════════
def build_qc_whiteboard_png():
    out = os.path.join(IMG_DIR, "qc_whiteboard.png")
    W, H = 1280, 900
    img = Image.new("RGB", (W, H), (236, 234, 220))
    d = ImageDraw.Draw(img)

    # whiteboard frame
    d.rectangle([40, 40, W-40, H-40], fill=(254, 254, 250), outline=(120, 113, 108), width=6)

    title = _font(40, bold=True)
    head = _font(28, bold=True)
    body = _font(22)
    small = _font(18)

    d.text((90, 80), "Trailspire 40L — Pre-Ship QC", fill=(15, 23, 42), font=title)
    d.text((90, 130), "Saigon TexCo · whiteboard, May 2026 (Marco Ferraro)", fill=(71, 85, 105), font=small)

    items = [
        ("1. Zipper test x10 cycles (main + 2 side pockets)", "YKK Vietnam ONLY — no Wenzhou substitutes"),
        ("2. Stitching density audit — 8 stitches / inch min", "Reject if any seam <7 spi"),
        ("3. Logo placement — within 2mm of tech-pack mark", "Use the placement template, not eyeballed"),
        ("4. Lot tag printed + sewn (inner seam, bottom-left)", "Lot # MUST match the PO; pink-ink lots = double-check"),
        ("5. Hardware torque check (all buckles + sternum)", ""),
        ("6. Water-bottle pocket elastic snap-back", ""),
        ("7. Pack the desiccant + lot card before sealing", ""),
    ]
    y = 200
    for item, sub in items:
        d.text((110, y), item, fill=(30, 41, 59), font=body)
        if sub:
            d.text((140, y + 32), "→ " + sub, fill=(159, 18, 57), font=small)
            y += 78
        else:
            y += 56

    # corner annotations
    d.text((90, H - 130), "⚠ Lot 24-A-118 → DEFECTIVE — DO NOT RESTOCK", fill=(190, 18, 60), font=head)
    d.text((90, H - 90), "(see #cs-ops Slack pin · Cara · 4 May)", fill=(71, 85, 105), font=small)

    # sticky note
    d.rectangle([W - 380, H - 280, W - 80, H - 80], fill=(254, 240, 138), outline=(133, 77, 14))
    d.text((W - 365, H - 270), "Sticky:", fill=(67, 39, 7), font=head)
    d.text((W - 365, H - 230), "Send 5 random units / lot\nto Marco for spot QC\nbefore wholesale ships.",
           fill=(67, 39, 7), font=body)

    img.save(out, "PNG", optimize=True)
    print(f"  ✓ {out}")


# ══════════════════════════════════════════════════════════════════════════════
# IMAGE 3 — Org chart screenshot
# ══════════════════════════════════════════════════════════════════════════════
def build_org_chart_png():
    out = os.path.join(IMG_DIR, "org_chart.png")
    W, H = 1500, 950
    img = Image.new("RGB", (W, H), (248, 250, 252))
    d = ImageDraw.Draw(img)

    title = _font(28, bold=True)
    sub = _font(15)
    d.text((40, 28), "Helix Outdoor — Org Chart", fill=(15, 23, 42), font=title)
    d.text((40, 70), "Generated from Notion · April 28, 2026 · 78 employees", fill=(100, 116, 139), font=sub)

    # CEO at top
    _box(d, (650, 110, 880, 200), "Diego Marin", "Founder / CEO",
         fill=(255, 255, 255), border=(15, 23, 42))

    # second row execs
    row2_y = 280
    row2 = [
        (80, 380, "Sai Krishnan", "CTO"),
        (440, 740, "Bruno Castelli", "Head of Finance"),
        (790, 1090, "VP Operations", "VACANT — Maya Lin departed Apr 12"),
        (1140, 1440, "Camille Rousseau", "Head of Customer Care"),
    ]
    for x1, x2, name, role in row2:
        is_vacant = "VACANT" in role
        _box(d, (x1, row2_y, x2, row2_y + 90), name, role,
             fill=(254, 226, 226) if is_vacant else (255, 255, 255),
             border=(127, 29, 29) if is_vacant else (15, 23, 42))
        # connect to CEO
        _arrow(d, (765, 200), ((x1 + x2) // 2, row2_y))

    # marketing + legal as advisors
    _box(d, (1100, 110, 1430, 200), "Lena Park / Priya Iyer",
         "Marketing · Legal (fractional)",
         fill=(241, 245, 249), border=(71, 85, 105))
    d.line([(880, 155), (1100, 155)], fill=(148, 163, 184), width=2)

    # third row reports
    row3_y = 480
    reports = [
        # Engineering
        (60, 220, row3_y, "Jin Ahn", "Senior Backend"),
        (60, 220, row3_y + 110, "Hira Qureshi", "Senior Frontend"),
        (60, 220, row3_y + 220, "Wei Zhao", "ML / Data"),
        (60, 220, row3_y + 330, "Tomas Becker", "DevOps"),
        # Finance
        (450, 600, row3_y, "Liam Walsh", "AP / AR"),
        (450, 600, row3_y + 110, "Renata Costa", "Controller (frac.)"),
        # Ops (under interim coverage)
        (790, 940, row3_y, "Marco Ferraro", "Acting Ops Lead"),
        (790, 940, row3_y + 110, "Nia Okafor", "Logistics Coordinator"),
        (790, 940, row3_y + 220, "Hugo Sanchez", "Inventory Planner"),
        # CX
        (1140, 1310, row3_y, "Aria Velasquez", "CX-DTC Lead"),
        (1140, 1310, row3_y + 110, "Cara Bennett", "CX-B2B Lead"),
        (1140, 1310, row3_y + 220, "Petra Olsson", "CX Specialist"),
        (1140, 1310, row3_y + 330, "Idris Khan", "CX Specialist"),
    ]
    for x1, x2, y, name, role in reports:
        _box(d, (x1, y, x2, y + 80), name, role,
             fill=(255, 255, 255), border=(71, 85, 105))

    # Legend
    legend_font = _font(13)
    d.text((40, H - 60),
           "Solid border = active role · red = vacant · italic role = fractional. "
           "Dashed exec line = advisory rather than direct report.",
           fill=(71, 85, 105), font=legend_font)

    img.save(out, "PNG", optimize=True)
    print(f"  ✓ {out}")


# ══════════════════════════════════════════════════════════════════════════════
def main():
    print("Building Helix Outdoor demo data...")
    print(" PDFs:")
    build_returns_policy()
    build_msa()
    build_vendor_risk()
    print(" DOCX:")
    build_security_policy()
    build_sla_matrix()
    print(" Images:")
    build_architecture_png()
    build_qc_whiteboard_png()
    build_org_chart_png()
    print("Done.")


if __name__ == "__main__":
    main()
